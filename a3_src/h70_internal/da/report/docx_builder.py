# -*- coding: utf-8 -*-
"""
Module for the generation of docx format documents.

---
type:
    python_module

validation_level:
    v00_minimum

protection:
    k00_public

copyright:
    "Copyright 2016 High Integrity Artificial Intelligence Systems"

license:
    "Licensed under the Apache License, Version 2.0 (the License);
    you may not use this file except in compliance with the License.
    You may obtain a copy of the License at

        http://www.apache.org/licenses/LICENSE-2.0

    Unless required by applicable law or agreed to in writing, software
    distributed under the License is distributed on an AS IS BASIS,
    WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
    See the License for the specific language governing permissions and
    limitations under the License."
...
"""


import os.path

import docx
import docx.oxml.shared
import lxml.html

import da.util
from . import docx_styles


ALIGN_LEFT   = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT    # pylint:disable=E1101
ALIGN_CENTRE = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # pylint:disable=E1101
PAGE_BREAK   = docx.enum.text.WD_BREAK.PAGE              # pylint:disable=E1101
NEW_PAGE     = docx.enum.section.WD_SECTION.NEW_PAGE     # pylint:disable=E1101


# -----------------------------------------------------------------------------
def build(doc_data, section_list, filepath):
    """
    Build and save the specified document.

    """
    document = docx.Document()
    docx_styles.set_styles(document)

    # _set_core_properties(doc_data, document)
    _add_title_section(document, doc_data['_metadata'])
    _add_toc_section(document)

    # print('#' * 80  )
    for section in section_list:
    #     root = lxml.html.fragment_fromstring(section['html'][0])

    #     if root.text:
    #         print('ROOT TEXT = ' + root.text)

    #     # print(repr(section['html'][0]))
    #     for child in root:
    #         if child.tail:
    #             print('CHILD TAIL = ' + child.tail)

        if section['level'] == 1:
            _add_new_page_section(document)

        if len(section['para']) > 0:
            _add_content_section(document, section)

    # Save the document.
    da.util.ensure_dir_exists(os.path.dirname(filepath))
    document.save(filepath)


# -----------------------------------------------------------------------------
def _add_content_section(document, section):
    """
    Add a content 'section' to the document.

    Note: Currently NOT implemented with a proper docx section, but still a
    set of paragraphs with a heading.

    """
    title_text  = '{num}\t{title}'.format(
                                    num   = section['num'],
                                    title = section['title'])

    title_style = document.styles['Heading {level}'.format(
                                    level = section['level'])]

    title_para  = document.add_paragraph(
                                    title_text,
                                    style = title_style)

    _suppress_line_numbers(title_para)

    # TODO: ITERATE OVER HTML 'RUNS' AND CONVERT TO DOCX XML RUNS

    for para in section['para']:

        if section['type'] == '_req':
            style = document.styles['Instructions']
        else:
            style = document.styles['Body Text']

        if para is not None:

            # TODO: RENDER MARKDOWN TO HTML THEN PARSE HTML

            document.add_paragraph(para, style = style)


# -----------------------------------------------------------------------------
def _add_title_section(document, doc_metadata):
    """
    Add a title page to the document.

    """
    section = document.sections[-1]

    # The title page has no header or footer and has a lowered top margin
    # to place the title about three quarters of the way up the page.
    docx_styles.set_default_margins(section)
    section.header_distance = docx.shared.Cm(0.0)
    section.footer_distance = docx.shared.Cm(0.0)
    section.top_margin      = docx.shared.Cm(6.0)

    # The document title specifies the document type and the name of
    # the system of interest.
    para_title = document.add_paragraph(
        '{document_type} ({document_type_acronym})\n'
        'for the\n'
        '{system_of_interest_name}'.format(
            document_type           = doc_metadata['document_type'],
            document_type_acronym   = doc_metadata['document_type_acronym'],
            system_of_interest_name = doc_metadata['system_of_interest_name']),
        style = document.styles['Title'])
    _suppress_line_numbers(para_title)

    # The document identification and metadata box contains the usual
    # identification and traceability information.
    para_meta = document.add_paragraph(
        'Document #\t\t{document_id}\n'
        'Compiled on:\t\t{compilation_date}\n'
        'Timebox #\t\t{timebox_id}\n'
        'Configuration #\t\t{configuration_id}\n'
        '\n'
        'Stage:\t\t\t{lifecycle_stage}\n'
        'Protection:\t\t{protective_marking}\n'
        '\n'
        'Contact person:\t\t{contact_person}\n'
        'Contact address:\t{contact_email}'.format(
            document_id         = doc_metadata['document_id'],
            compilation_date    = doc_metadata['compilation_date'],
            timebox_id          = doc_metadata['timebox_id'],
            configuration_id    = doc_metadata['configuration_id'],
            lifecycle_stage     = doc_metadata['lifecycle_stage'],
            protective_marking  = doc_metadata['protective_marking'],
            contact_person      = doc_metadata['contact_person'],
            contact_email       = doc_metadata['contact_email']),
        style = document.styles['Body Text'])

    # We position the document identification at the bottom right hand
    # corner of the front page analogously to the location of the title
    # block of a technical drawing.
    # http://officeopenxml.com/WPparagraph-textFrames.php
    qualname = docx.oxml.shared.qn
    frame    = docx.oxml.shared.OxmlElement('w:framePr')
    frame.set(qualname('w:AnchorLock'),   'true')
    frame.set(qualname('w:w'),            '5200')
    frame.set(qualname('w:h'),            '5200')
    frame.set(qualname('w:wrap'),         'auto')
    frame.set(qualname('w:vAnchor'),      'margin')
    frame.set(qualname('w:hAnchor'),      'margin')
    frame.set(qualname('w:xAlign'),       'right')
    frame.set(qualname('w:yAlign'),       'bottom')
    para_props = para_meta._element.pPr                 # pylint: disable=W0212
    para_props.insert_element_before(
                    frame, 'w:keepNext', 'w:keepLines', 'w:spacing', 'w:jc')
    _suppress_line_numbers(para_meta)

    _add_intentionally_blank_page(document)

    # # TODO: Signature page
    # para_title = document.add_paragraph(
    #     'Approval',
    #     style = document.styles['Heading 1'])
    # _suppress_line_numbers(para_title)

    # _add_intentionally_blank_page(document)

    # # TODO: Preface
    # para_title = document.add_paragraph(
    #     'Preface',
    #     style = document.styles['Heading 1'])
    # _suppress_line_numbers(para_title)

    # _add_intentionally_blank_page(document)

    # # TODO: Abstract
    # para_title = document.add_paragraph(
    #     'Abstract',
    #     style = document.styles['Heading 1'])
    # _suppress_line_numbers(para_title)

    # _add_intentionally_blank_page(document)

    # # TODO: Change information page
    # para_title = document.add_paragraph(
    #     'Change information',
    #     style = document.styles['Heading 1'])
    # _suppress_line_numbers(para_title)

    # _add_intentionally_blank_page(document)


# -----------------------------------------------------------------------------
def _add_intentionally_blank_page(document):
    """
    Add a blank page with the message: 'This page intentionally left blank'.

    """
    _add_paragraph_with_page_break(document)
    para_blank = document.add_paragraph('This page intentionally left blank',
                                        style = document.styles['Body Text'])
    para_blank.alignment = ALIGN_CENTRE
    _suppress_line_numbers(para_blank)

    _add_paragraph_with_page_break(document)


# -----------------------------------------------------------------------------
def _add_paragraph_with_page_break(document):
    """
    Add a page break.

    I needed to create this function so that I could suppress line numbers
    on the (empty) paragraph containing the page break.

    """
    para = document.add_paragraph('', style = document.styles['Body Text'])
    _suppress_line_numbers(para)
    para.add_run().add_break(PAGE_BREAK)


# -----------------------------------------------------------------------------
def _add_toc_section(document):
    """
    Add a table of contents section to the document.

    """
    document.add_section(NEW_PAGE)
    docx_styles.set_default_margins(document.sections[-1])
    toc_para = document.add_paragraph('Table of Contents\n ',
                          style = document.styles['Heading 1'])
    toc_run  = toc_para.add_run()
    r_elem   = toc_run._r                               # pylint: disable=W0212

    # begin
    fld_begin = docx.oxml.shared.OxmlElement('w:fldChar')
    fld_begin.set(docx.oxml.shared.qn('w:fldCharType'), 'begin')
    r_elem.append(fld_begin)

    # instrText
    instr_txt = docx.oxml.shared.OxmlElement('w:instrText')
    instr_txt.set(docx.oxml.shared.qn('xml:space'), 'preserve')
    instr_txt.text  = r'TOC \o "1-2" \h \z \u'
    r_elem.append(instr_txt)

    # separate
    fld_sep = docx.oxml.shared.OxmlElement('w:fldChar')
    fld_sep.set(docx.oxml.shared.qn('w:fldCharType'), 'separate')
    fld_ctrl = docx.oxml.shared.OxmlElement('w:t')
    fld_ctrl.text = "Right-click to update field."
    fld_sep.append(fld_ctrl)
    r_elem.append(fld_sep)

    # end
    fld_end = docx.oxml.shared.OxmlElement('w:fldChar')
    fld_end.set(docx.oxml.shared.qn('w:fldCharType'), 'end')
    r_elem.append(fld_end)

    _suppress_line_numbers(toc_para)

    # Trigger auto-update of the TOC next time word is opened.
    update_fields = docx.oxml.shared.OxmlElement('w:updateFields')
    document.settings.element.append(update_fields)


# -----------------------------------------------------------------------------
def _suppress_line_numbers(para):
    """
    Set the suppressLineNumbers attribute on the specified paragraph.

    """
    para_props = para._element.pPr                      # pylint: disable=W0212
    para_props.append(docx.oxml.shared.OxmlElement('w:suppressLineNumbers'))


# -----------------------------------------------------------------------------
def _add_new_page_section(document):
    """
    Add a new page section to the document.

    """
    section = document.add_section(NEW_PAGE)
    docx_styles.set_default_margins(section)

    # Set line numbering
    qualname    = docx.oxml.shared.qn
    ln_num_type = docx.oxml.shared.OxmlElement('w:lnNumType')
    ln_num_type.set(qualname('w:start'),      '0')            # Starting value.
    ln_num_type.set(qualname('w:countBy'),    '2')            # Increment.
    ln_num_type.set(qualname('w:restart'),    'newPage')      # Restart policy.
    ln_num_type.set(qualname('w:distance'),   '100')
    sect_props = section._sectPr                        # pylint: disable=W0212
    sect_props.insert_element_before(
        ln_num_type, 'w:pgNumType', 'w:cols', 'w:formProt', 'w:vAlign',
        'w:noEndnote', 'w:titlePg', 'w:textDirection', 'w:bidi',
        'w:rtlGutter', 'w:docGrid', 'w:printerSettings')

    # <w:lnNumType w:countBy="3" w:start="1" w:restart="newSection"/>
    #
    # countBy   Specifies the line number increments to be displayed.
    #           So, e.g., a value of 5 would specify that a page number
    #           appear at every fifth page.
    #
    # distance  Specifies the distance between text and line numbering,
    #           in twips (i.e., 1/1440th of an inch).
    #
    # restart   Specifies when numbering should be reset to the line
    #           number specified in the start attribute. Possible values
    #           are:
    #
    #           continuous  numbering should continue from the number
    #                       of the previous section.
    #           newPage     numbering should restart when a new page
    #                       is displayed. This is the default value.
    #           newSection  number should restart when the section begins.
    #
    # start     Specifies the starting value for the numbering.


# -----------------------------------------------------------------------------
# def _set_core_properties(doc_data, document):
#    """
#    Set the core properties of the document.
#
#    The core properties are author, category, comments, content_status,
#    created, identifier, keywords, language, last_modified_by, last_printed,
#    modified, revision, subject, title, and version.
#
#    """

    # Author.
    # string – An entity primarily responsible for making the content of the
    # resource.
    # document.core_properties.author = 'author'

    # string – A categorization of the content of this package. Example values
    # might include: Resume, Letter, Financial Forecast, Proposal, or
    # Technical Presentation.
    # document.core_properties.category = 'category'

    # string – An account of the content of the resource.
    # document.core_properties.comments = 'comments'

    # string – completion status of the document, e.g. ‘draft’
    # document.core_properties.content_status = 'content_status'

    # datetime – time of intial creation of the document
    # document.core_properties.created = 'created'

    # string – An unambiguous reference to the resource within a given
    # context, e.g. ISBN.
    # document.core_properties.identifier = 'identifier'

    # string – descriptive words or short phrases likely to be used as search
    # terms for this document
    # document.core_properties.keywords = 'keywords'

    # string – language the document is written in
    # document.core_properties.language = 'language'

    # string – name or other identifier (such as email address) of person
    # who last modified the document
    # document.core_properties.last_modified_by = 'last_modified_by'

    # datetime – time the document was last printed
    # document.core_properties.last_printed = 'last_printed'

    # datetime – time the document was last modified
    # document.core_properties.modified = 'modified'

    # int – number of this revision, incremented by Word each time the
    # document is saved. Note however python-docx does not automatically
    # increment the revision number when it saves a document.
    # document.core_properties.revision = 0

    # string – The topic of the content of the resource.
    # document.core_properties.subject = 'subject'

    # string – The name given to the resource.
    # document.core_properties.title = 'title'

    # string – free-form version string
    # document.core_properties.version =
